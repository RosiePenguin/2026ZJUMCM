from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
RESULTS = ROOT / "results"


def set_page_margins(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def start_page_numbering(section, start: int = 1) -> None:
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    section._sectPr.append(pg_num_type)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for section in doc.sections:
        set_page_margins(section)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(14 if level == 1 else 12)


def add_paragraph(doc: Document, text: str, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def build_doc() -> Document:
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("“浙超”分组方案、抽签机制、赛点选择与赛制优化的一体化研究")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)
    add_paragraph(doc, "此页用于替换为学校统一封面。", center=True)

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(sec)
    start_page_numbering(sec, 1)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    add_heading(doc, "摘要")
    add_paragraph(doc, "针对“浙超”联赛 64 支队伍的小组分组、抽签公平性、比赛地点选择及赛制优化问题，本文构建了一套“分组生成—公平抽签—赛点配置—赛制优化”的一体化建模框架，并将题目四问统一为一条可复现、可解释、可扩展的决策链。问题 1 中，本文建立带硬约束与软约束的组合优化模型，并设计“市级队先随机落位、县级队按约束紧度逐步安置”的随机启发式算法，在 500 次搜索中得到 450 个可行方案；最优方案实现同市县级队同组冲突对数为 0、平均城市来源熵为 2.0000、各组总实力标准差为 6.2099。问题 2 中，本文提出分阶段等概率抽签流程，并通过 10000 次蒙特卡洛模拟和卡方检验验证其公平性，结果显示平均 p 值为 0.5418，无队伍出现显著偏差。问题 3 中，本文基于浙江省 64 个参赛地区的人口、GDP、交通、体育基础、铁路与机场可达性等数据，建立熵权多目标离散设施选址模型，推荐三门县、诸暨市、仙居县、丽水市、义乌市、东阳市、磐安县和金华市 8 个赛点，平均分配距离为 127.83 km，最大距离为 214.82 km。问题 4 中，本文利用熵权合成的实力代理值与带疲劳项的比赛仿真模型，对 7 类赛制进行了 20000 次蒙特卡洛比较。结果表明，seeded_reseeded 方案在优先级加权 TOPSIS 下得分最高，为 0.822888；相较于当前随机赛制，该方案将前 4 强队在 32 强首轮出局概率由 0.2831 降至 0.2551，并将前 8 强队进入八强的概率由 0.3817 提高到 0.3968。进一步地，本文从分组多方案比较、抽签统计公平性、赛点多目标权衡和赛制疲劳传递建模四个方面对模型进行了系统改进，并通过敏感性分析验证了主要结论的稳健性。")
    add_paragraph(doc, "关键词：分组优化；公平抽签；设施选址；蒙特卡洛模拟；熵权法；TOPSIS")

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(doc.sections[-1])
    footer2 = doc.sections[-1].footer.paragraphs[0]
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer2)

    add_heading(doc, "1 问题重述")
    add_paragraph(doc, "“浙超”联赛共有 64 支队伍参赛，其中包括 11 支市级队、20 支县级市队和 33 支县队。比赛分为两个阶段：第一阶段为 16 个小组、每组 4 队的单循环小组赛，小组前两名晋级；第二阶段为 32 强淘汰赛，经过 5 轮决出冠军。本文按“分组—抽签—赛点—赛制”的链式结构统一求解。")

    add_heading(doc, "2 问题分析与模型假设")
    add_paragraph(doc, "问题 1 属于带硬约束与软约束的组合优化问题，问题 2 属于带约束的随机公平性问题，问题 3 属于多目标离散设施选址问题，问题 4 属于基于蒙特卡洛仿真的赛制评价问题。")
    add_bullets(doc, [
        "64 支队伍全部按题意参赛，不考虑中途退赛。",
        "行政隶属关系固定，市级队与其代管县级队的回避关系不发生变化。",
        "缺乏完整历史战绩时，用地区综合指标构造球队实力代理值。",
        "抽签设备与程序在技术上等概率执行，不存在人为操控。",
        "比赛出行成本主要由地理距离和交通可达性刻画。",
    ])

    add_heading(doc, "3 问题 1：分组方案生成与评价")
    add_paragraph(doc, "本文将分组要求拆分为硬约束与软约束，采用“市级队先随机落位、县级队按约束紧度逐步放置”的随机启发式算法。程序共搜索 500 次，得到 450 个可行方案。")
    add_table(doc, ["指标", "数值"], [["可行方案数", "450"], ["同市县级队同组冲突对数", "0.0"], ["平均城市熵", "2.0000"], ["各组总实力标准差", "6.2099"], ["综合得分", "187.5803"]])
    add_paragraph(doc, "该方案不仅完全满足全部硬约束，而且在县级队分散性和小组实力均衡性方面表现良好。")

    add_heading(doc, "4 问题 2：公平抽签方案设计与验证")
    add_paragraph(doc, "本文采用“市级队先抽、县级队后抽”的分阶段流程。市级队首先随机进入 16 个小组中的 11 个不同小组；县级队随后在实时更新的合法小组集合中等概率抽签。通过 10000 次蒙特卡洛模拟，64 支队伍的平均 p 值为 0.5418，无队伍出现 p<0.01 的显著偏差。")
    add_figure(doc, RESULTS / "problem2" / "fairness_analysis.png", "图 1 抽签公平性统计图")
    add_figure(doc, RESULTS / "problem2" / "draw_example.png", "图 2 单次抽签示例")

    add_heading(doc, "5 问题 3：比赛地点选择模型")
    add_paragraph(doc, "本文利用 64 个参赛地区的人口、GDP、交通、体育基础、铁路可达性、机场可达性及代表性场馆容量等信息，构建赛事影响力指标与承办能力指标，并建立多目标离散设施选址模型。")
    add_table(doc, ["赛点", "承办小组", "平均距离/km", "最大距离/km"], [["三门县", "G01+G02", "163.65", "209.14"], ["诸暨市", "G03+G12", "120.81", "157.00"], ["仙居县", "G04+G09", "94.14", "145.57"], ["丽水市", "G05+G15", "105.25", "153.06"], ["义乌市", "G06+G14", "136.61", "176.51"], ["东阳市", "G07+G10", "106.19", "150.84"], ["磐安县", "G08+G11", "162.17", "214.82"], ["金华市", "G13+G16", "133.80", "181.39"]])
    add_paragraph(doc, "整体方案的平均分配距离为 127.83 km，最大分配距离为 214.82 km，且不存在主场惩罚触发情形。")
    add_figure(doc, RESULTS / "problem34" / "problem3_typhoon_sensitivity.png", "图 3 赛点方案对极端天气权重的敏感性")

    add_heading(doc, "6 问题 4：赛制优化模型")
    add_paragraph(doc, "本文基于熵权法合成的实力代理值和带疲劳项的比赛仿真模型，对 7 类候选赛制进行比较。结果显示，seeded_reseeded 方案在综合评价中最优。")
    add_table(doc, ["赛制", "前8进八强率", "前4首轮出局率", "TOPSIS"], [["current_random_draw", "0.3817", "0.2831", "0.258954"], ["group_rank_seeded", "0.3894", "0.2575", "0.662609"], ["seeded_reseeded", "0.3968", "0.2551", "0.822888"], ["seeded_reseeded_rest", "0.3931", "0.2591", "0.706284"]])
    add_figure(doc, RESULTS / "problem34" / "problem4_format_comparison.png", "图 4 各候选赛制的综合比较")

    add_heading(doc, "7 模型改进与深化")
    add_bullets(doc, [
        "将问题 1 从单一可行性搜索提升为多方案生成与多指标排序。",
        "将问题 2 从规则描述提升为统计公平性检验。",
        "将问题 3 从单一距离优化提升为多目标离散设施选址。",
        "将问题 4 从静态赛制比较提升为带疲劳传递的动态仿真。",
        "对问题 3 和问题 4 的关键权重与参数进行了敏感性分析。",
    ])

    add_heading(doc, "8 结论")
    add_bullets(doc, [
        "得到了满足全部硬约束且软约束冲突为 0 的优良分组方案。",
        "抽签机制在统计意义上公平透明。",
        "推荐了 8 个兼顾便利与赛事影响力的比赛地点。",
        "推荐采用“小组第一对小组第二 + 后续复排”的 seeded_reseeded 赛制。",
    ])

    add_heading(doc, "9 附录与支撑材料说明")
    add_paragraph(doc, "根据竞赛要求，论文附录与支撑材料应完整包含全部可运行源程序、实际使用的软件名称和命令、自主查阅使用的数据与资料，而不应重复放入赛题已直接提供的数据文件。")
    add_bullets(doc, [
        "源程序代码：src/problem1、src/problem2、src/problem34 下的全部 Python 文件。",
        "运行脚本与命令：run_problem1.py、run_problem2.py、run_problem34.py、clean_problem3_manual_sources.py、generate_b_paper_docx.py。",
        "自主查阅并整理的数据：region_metrics.csv、problem3_missing_fields_lookup.csv 及人工收集的 Excel 数据源。",
        "结果支撑文件：problem1、problem2、problem34 对应的关键 csv、txt 与 png 输出文件。",
        "AI 使用说明：docs/AI工具使用说明.md。",
    ])

    add_heading(doc, "附录 A AI 工具使用说明")
    add_bullets(doc, [
        "工具名称：OpenAI Codex",
        "模型：GPT-5 系列编码代理",
        "开发机构：OpenAI",
        "使用日期：2026 年 5 月 14 日至 2026 年 5 月 17 日",
        "用途：论文结构梳理、结果文字转写、图表说明和文档排版辅助。",
    ])

    return doc


def main() -> None:
    out = DOCS / "B题论文终稿.docx"
    doc = build_doc()
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
